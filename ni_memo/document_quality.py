"""Structural DOCX fidelity checks independent from fact-quality gates."""
from dataclasses import asdict, dataclass
import filecmp
from pathlib import Path
import zipfile
from xml.etree import ElementTree

from docx import Document


W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


@dataclass(frozen=True)
class DocumentProfile:
    paragraph_characters: int
    table_count: int
    inline_shape_count: int
    media_count: int
    section_count: int
    footnote_count: int
    rendered_page_break_count: int

    def to_dict(self):
        return asdict(self)


@dataclass(frozen=True)
class FidelityReport:
    status: str
    byte_identical: bool
    candidate: DocumentProfile
    reference: DocumentProfile
    ratios: dict[str, float]
    failed_metrics: tuple[str, ...]

    def to_dict(self):
        return {
            "status": self.status,
            "byte_identical": self.byte_identical,
            "candidate": self.candidate.to_dict(),
            "reference": self.reference.to_dict(),
            "ratios": dict(self.ratios),
            "failed_metrics": list(self.failed_metrics),
        }


@dataclass(frozen=True)
class CompositionReport:
    status: str
    paragraph_count: int
    paragraph_characters: int
    long_paragraph_count: int
    table_characters: int
    body_share: float
    failed_metrics: tuple[str, ...]

    def to_dict(self):
        data = asdict(self)
        data["failed_metrics"] = list(self.failed_metrics)
        return data


@dataclass(frozen=True)
class FormalTemplateReport:
    status: str
    reference_name: str
    content_identity_required: bool
    checks: dict[str, bool]
    failed_metrics: tuple[str, ...]
    missing_sections: tuple[str, ...] = ()
    missing_slots: tuple[str, ...] = ()
    out_of_order_slots: tuple[str, ...] = ()
    section_layout_match: bool = False
    style_checks: dict[str, bool] = None
    table_checks: dict[str, bool] = None

    def to_dict(self):
        return {
            "status": self.status,
            "reference_name": self.reference_name,
            "content_identity_required": self.content_identity_required,
            "checks": dict(self.checks),
            "failed_metrics": list(self.failed_metrics),
            "missing_sections": list(self.missing_sections),
            "missing_slots": list(self.missing_slots),
            "out_of_order_slots": list(self.out_of_order_slots),
            "section_layout_match": self.section_layout_match,
            "style_checks": dict(self.style_checks or {}),
            "table_checks": dict(self.table_checks or {}),
        }


def profile_docx(path):
    source = Path(path)
    document = Document(source)
    with zipfile.ZipFile(source) as package:
        names = package.namelist()
        media_count = sum(name.startswith("word/media/") and not name.endswith("/") for name in names)
        footnote_count = _footnote_count(package)
        page_breaks = _page_break_count(package)
    return DocumentProfile(
        paragraph_characters=sum(len(paragraph.text) for paragraph in document.paragraphs),
        table_count=len(document.tables),
        inline_shape_count=len(document.inline_shapes),
        media_count=media_count,
        section_count=len(document.sections),
        footnote_count=footnote_count,
        rendered_page_break_count=page_breaks,
    )


def evaluate_document_composition(path, min_characters=10_000, min_long_paragraphs=20,
                                  min_body_share=0.35):
    """Enforce the memo's readable-analysis floor independently from fact quality."""
    document = Document(path)
    paragraphs = tuple(
        paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()
    )
    paragraph_characters = sum(len(text) for text in paragraphs)
    table_characters = sum(
        len(cell.text.strip())
        for table in document.tables for row in table.rows for cell in row.cells
    )
    total_characters = paragraph_characters + table_characters
    body_share = paragraph_characters / total_characters if total_characters else 0.0
    metrics = {
        "paragraph_characters": paragraph_characters >= min_characters,
        "long_paragraph_count": sum(len(text) >= 100 for text in paragraphs) >= min_long_paragraphs,
        "body_share": body_share >= min_body_share,
    }
    failed = tuple(key for key, passed in metrics.items() if not passed)
    return CompositionReport(
        status="pass" if not failed else "fail",
        paragraph_count=len(paragraphs),
        paragraph_characters=paragraph_characters,
        long_paragraph_count=sum(len(text) >= 100 for text in paragraphs),
        table_characters=table_characters,
        body_share=round(body_share, 6),
        failed_metrics=failed,
    )


def compare_document_fidelity(candidate, reference):
    candidate_path, reference_path = Path(candidate), Path(reference)
    candidate_profile = profile_docx(candidate_path)
    reference_profile = profile_docx(reference_path)
    fields = tuple(DocumentProfile.__dataclass_fields__)
    ratios = {
        field: _ratio(getattr(candidate_profile, field), getattr(reference_profile, field))
        for field in fields
    }
    failed = tuple(
        field for field in fields
        if getattr(candidate_profile, field) != getattr(reference_profile, field)
    )
    identical = filecmp.cmp(candidate_path, reference_path, shallow=False)
    return FidelityReport(
        status="pass" if not failed else "fail",
        byte_identical=identical,
        candidate=candidate_profile,
        reference=reference_profile,
        ratios=ratios,
        failed_metrics=failed,
    )


def compare_formal_template(candidate, reference, schema):
    """Compare the reusable editorial/layout contract, never project facts."""
    candidate_path, reference_path = Path(candidate), Path(reference)
    candidate_doc, reference_doc = Document(candidate_path), Document(reference_path)
    expected_sections = tuple(section.title for section in schema.sections)
    expected_slots = tuple(block.title for block in schema.blocks if block.required)
    candidate_sections = _heading_contract(candidate_doc, "Heading 1", expected_sections)
    candidate_slots = _heading_contract(candidate_doc, "Heading 2", expected_slots)
    missing_sections = tuple(title for title in expected_sections if title not in candidate_sections)
    missing_slots = tuple(title for title in expected_slots if title not in candidate_slots)
    section_order = candidate_sections == expected_sections
    slot_order = candidate_slots == expected_slots
    out_of_order_slots = () if slot_order else candidate_slots

    expected_layouts = _expected_section_layouts(schema)
    candidate_layouts = tuple(_section_signature(section) for section in candidate_doc.sections)
    reference_layouts = tuple(_section_signature(section) for section in reference_doc.sections)
    section_layouts = candidate_layouts == expected_layouts == reference_layouts

    style_checks = {
        name.lower().replace(" ", "_"): _style_signature(candidate_doc, name)
        == _style_signature(reference_doc, name)
        for name in ("Normal", "Heading 1", "Heading 2", "Table Grid")
    }
    table_checks = _table_contract(candidate_doc)
    with zipfile.ZipFile(candidate_path) as package:
        document_xml = package.read("word/document.xml").decode("utf-8")
        footer_xml = "\n".join(
            package.read(name).decode("utf-8")
            for name in package.namelist() if name.startswith("word/footer") and name.endswith(".xml")
        )
    checks = {
        "section_layouts": section_layouts,
        "section_order": section_order and not missing_sections,
        "required_slots": not missing_slots,
        "slot_order": slot_order and not missing_slots,
        **style_checks,
        **table_checks,
        "toc_field": " TOC " in document_xml,
        "page_field": " PAGE " in footer_xml,
    }
    failed = tuple(key for key, passed in checks.items() if not passed)
    return FormalTemplateReport(
        status="pass" if not failed else "fail",
        reference_name=reference_path.name,
        content_identity_required=False,
        checks=checks,
        failed_metrics=failed,
        missing_sections=missing_sections,
        missing_slots=missing_slots,
        out_of_order_slots=out_of_order_slots,
        section_layout_match=section_layouts,
        style_checks=style_checks,
        table_checks=table_checks,
    )


def _style_signature(document, name, include_bold=True):
    style = document.styles[name]
    fonts = style._element.rPr.rFonts if style._element.rPr is not None else None
    east_asia = fonts.get(f"{{{W_NS}}}eastAsia") if fonts is not None else None
    signature = (
        style.font.name,
        east_asia,
        int(style.font.size) if style.font.size is not None else None,
    )
    return signature + ((style.font.bold,) if include_bold else ())


def _style_size(document, name):
    size = document.styles[name].font.size
    return int(size) if size is not None else None


def _heading_contract(document, style_name, expected_titles):
    values = []
    for paragraph in document.paragraphs:
        if paragraph.style.name != style_name:
            continue
        matches = [title for title in expected_titles if title in paragraph.text]
        if matches:
            values.append(max(matches, key=len))
    return tuple(values)


def _expected_section_layouts(schema):
    layouts = []
    for index, section in enumerate(schema.sections):
        if index == 0 or section.new_word_section:
            if section.orientation == "landscape":
                layouts.append((10692130, 7560310, 1141095, 1141095, 914400, 914400))
            else:
                layouts.append((7560310, 10692130, 914400, 914400, 1143000, 1143000))
    return tuple(layouts)


def _section_signature(section):
    return tuple(int(getattr(section, name)) for name in (
        "page_width", "page_height", "top_margin", "bottom_margin", "left_margin", "right_margin",
    ))


def _table_contract(document):
    tables = tuple(table for table in document.tables if len(table.rows) >= 2)
    if not tables:
        return {"table_headers": False, "table_rows": False}
    headers = all("w:tblHeader" in table.rows[0]._tr.xml for table in tables)
    rows = all(
        "w:cantSplit" in row._tr.xml
        for table in tables for row in table.rows[1:]
    )
    return {"table_headers": headers, "table_rows": rows}


def _ratio(candidate, reference):
    return 1.0 if reference == 0 and candidate == 0 else (candidate / reference if reference else 0.0)


def _footnote_count(package):
    if "word/footnotes.xml" not in package.namelist():
        return 0
    root = ElementTree.fromstring(package.read("word/footnotes.xml"))
    identifier = f"{{{W_NS}}}id"
    return sum(int(node.attrib.get(identifier, "-1")) >= 0 for node in root.findall(f"{{{W_NS}}}footnote"))


def _page_break_count(package):
    total = 0
    for name in ("word/document.xml",):
        if name not in package.namelist():
            continue
        root = ElementTree.fromstring(package.read(name))
        total += len(root.findall(f".//{{{W_NS}}}lastRenderedPageBreak"))
        type_key = f"{{{W_NS}}}type"
        total += sum(node.attrib.get(type_key) == "page" for node in root.findall(f".//{{{W_NS}}}br"))
    return total
