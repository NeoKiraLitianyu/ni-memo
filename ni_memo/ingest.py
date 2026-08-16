"""Deterministic multi-format evidence ingestion."""
import csv
import json
import re
from collections import defaultdict
from datetime import date, datetime
from pathlib import Path

from docx import Document
from openpyxl import load_workbook

from ni_memo.evidence import EvidenceItem, IngestResult

SUPPORTED = {".xlsx", ".docx", ".pdf", ".csv", ".json"}


def ingest_inputs(paths):
    files, skipped = _expand(paths)
    items, errors = [], []
    for path in files:
        try:
            items.extend(_LOADERS[path.suffix.lower()](path))
        except Exception as exc:
            errors.append(f"{path}: {type(exc).__name__}: {exc}")
    items.sort(key=lambda item: (item.source_uri.lower(), item.source_location, item.label))
    return IngestResult(tuple(items), tuple(errors), tuple(sorted(skipped)))


def _expand(paths):
    found, skipped = set(), set()
    for raw in paths:
        path = Path(raw).resolve()
        candidates = sorted(path.rglob("*")) if path.is_dir() else [path]
        for candidate in candidates:
            if not candidate.is_file():
                continue
            if candidate.suffix.lower() in SUPPORTED:
                found.add(candidate)
            else:
                skipped.add(str(candidate))
    return sorted(found, key=lambda path: str(path).lower()), skipped


def _csv(path):
    result = []
    with path.open("r", encoding="utf-8-sig", newline="") as handle:
        rows = list(csv.reader(handle))
    if not rows:
        return result
    headers = rows[0]
    for row_no, row in enumerate(rows[1:], start=2):
        row_label = row[0].strip() if row else ""
        for col_no, value in enumerate(row[1:], start=2):
            if value == "":
                continue
            header = headers[col_no - 1].strip() if col_no <= len(headers) else ""
            result.append(_item(path, "csv", f"row={row_no},col={col_no}", row_label or header,
                                _scalar(value), raw=value, context=(header,)))
    return result


def _json(path):
    data = json.loads(path.read_text(encoding="utf-8"))
    result = []

    def walk(value, parts):
        if isinstance(value, dict):
            for key in sorted(value):
                walk(value[key], parts + [str(key)])
        elif isinstance(value, list):
            for index, item in enumerate(value):
                walk(item, parts + [str(index)])
        else:
            location = "$" + "".join(f".{part}" if not part.isdigit() else f"[{part}]" for part in parts)
            result.append(_item(path, "json", location, parts[-1] if parts else "$", value,
                                raw=json.dumps(value, ensure_ascii=False), context=tuple(parts[:-1])))

    walk(data, [])
    return result


def _docx(path):
    doc, result = Document(path), []
    current_heading = ""
    for index, paragraph in enumerate(doc.paragraphs, start=1):
        text = paragraph.text.strip()
        if not text:
            continue
        if _is_heading(paragraph):
            current_heading = text
            continue
        label, value = _key_value(text)
        if label == text and value == text:
            label = current_heading or "段落"
        result.append(_item(path, "docx", f"paragraph={index}", label, value, raw=text,
                            context=(current_heading,) if current_heading else ()))
        result.extend(_docx_semantic_candidates(path, index, text, current_heading))
    for table_no, table in enumerate(doc.tables, start=1):
        for row_no, row in enumerate(table.rows, start=1):
            texts = [cell.text.strip() for cell in row.cells]
            if len(texts) >= 2 and texts[0] and texts[1]:
                result.append(_item(path, "docx", f"table={table_no},row={row_no},col=2",
                                    texts[0], _scalar(texts[1]), raw=" | ".join(texts)))
    return result


def _is_heading(paragraph):
    style_name = str(getattr(getattr(paragraph, "style", None), "name", "") or "").casefold()
    return "heading" in style_name or "标题" in style_name


def _docx_semantic_candidates(path, paragraph_index, text, heading):
    """Emit narrowly scoped facts derived directly from a source paragraph.

    These are candidates, not asserted truth: the exact paragraph and source text
    stay attached so reconciliation and acceptance gates can challenge them.
    """
    result = []
    company_pattern = (
        r"(?:^|[，。；：、“（(])"
        r"([\u4e00-\u9fffA-Za-z0-9（）()·&\-]{2,80}?"
        r"(?:股份有限公司|有限责任公司|有限公司))"
    )
    for match in re.finditer(company_pattern, text):
        candidate = match.group(1)
        if "子公司" in candidate:
            continue
        result.append(_item(
            path, "docx", f"paragraph={paragraph_index},entity=company",
            "公司名称候选", candidate, raw=text,
            context=(heading,) if heading else (),
        ))
    founded = re.search(r"(?:成立于|创立于)\s*((?:19|20)\d{2})\s*年?", text)
    if founded:
        result.append(_item(
            path, "docx", f"paragraph={paragraph_index},entity=founded",
            "成立时间候选", founded.group(1), raw=text,
            context=(heading,) if heading else (),
            coerce_value=False,
        ))
    return result


def _xlsx(path):
    formulas = load_workbook(path, data_only=False, read_only=False)
    values = load_workbook(path, data_only=True, read_only=False)
    result = []
    try:
        for sheet in formulas.worksheets:
            value_sheet = values[sheet.title]
            cells = _instantiated_cells(sheet)
            row_index, column_index = _sparse_indexes(cells)
            sheet_unit = _sheet_unit(cells)
            for cell in cells:
                raw = cell.value
                if isinstance(raw, bool):
                    continue
                cached = _cell_value(value_sheet, cell.row, cell.column)
                is_formula = cell.data_type == "f" or (isinstance(raw, str) and raw.startswith("="))
                formula_text = _formula_text(raw) if is_formula else None
                if isinstance(raw, str) and not is_formula:
                    label, value = _key_value(raw)
                    if label == raw and value == raw:
                        nearby_label = (_nearest_label(row_index, cell.row, cell.column)
                                        or _nearest_column_label(column_index, cell.row, cell.column))
                        if not nearby_label or nearby_label == raw:
                            continue
                        label, value = nearby_label, raw
                else:
                    label = _nearest_label(row_index, cell.row, cell.column)
                    value = cached if is_formula else raw
                context = tuple(filter(None, (
                    _section_title(row_index, cell.row, cell.column),
                    _nearest_header(column_index, cell.row, cell.column),
                    _parent_metric(row_index, cell.row, cell.column),
                )))
                unit = _cell_unit(cell, label, sheet_unit)
                result.append(_item(path, "xlsx", f"{sheet.title}!{cell.coordinate}", label, value, unit=unit,
                                    raw=formula_text or str(raw), formula=formula_text,
                                    cached=cached, context=context))
    finally:
        formulas.close()
        values.close()
    return result


def _pdf(path):
    from pypdf import PdfReader
    result = []
    for page_no, page in enumerate(PdfReader(str(path)).pages, start=1):
        text = page.extract_text() or ""
        for line_no, line in enumerate(text.splitlines(), start=1):
            line = line.strip()
            if not line:
                continue
            label, value = _key_value(line)
            result.append(_item(path, "pdf", f"page={page_no},line={line_no}", label, value, raw=line))
    return result


def _nearest_label(row_index, row, column):
    for col, value in reversed(row_index.get(row, ())):
        if col >= column:
            continue
        if column - col > 12:
            break
        if isinstance(value, str) and value.strip() and not value.startswith("="):
            return value.strip()
    return ""


def _section_title(row_index, row, column):
    """Find the nearest section/block title above this cell (e.g. 项目概览/退出假设).

    Summary matrices split the sheet into titled blocks (项目概览, 本次投资, 退出假设,
    ...). A cell deep in a block must carry its block title in context so guards
    can tell e.g. the 项目概览 投后估值 from the 退出假设 估值. A block title is a
    text in the label column whose row carries NO numeric value to its right —
    data-row labels (估值/股本/P/E...) always have a number in the same row, so
    they are never block titles. Column-header tokens (投前/投后/年份/审计) are
    also not block titles.
    """
    _metric_tokens = ("估值", "价格", "股本", "金额", "出资", "合计", "小计", "占比",
                      "比例", "倍数", "利率", "汇率", "成本", "收入", "毛利", "净利",
                      "费用", "现金", "资产", "负债", "权益", "总额", "规模", "%", "yoy")
    _header_tokens = ("投前", "投后", "审计", "单位", "币种", "预测", "实际")
    for candidate_row in range(row - 1, max(0, row - 8), -1):
        # 表头区的无数值文本是列头/表标题, 由 _nearest_header 负责; 区块标题
        # 只出现在数据区中间 (row > 4).
        if candidate_row <= 4:
            continue
        cells_in_row = sorted(row_index.get(candidate_row, ()))
        if not cells_in_row:
            continue
        label_col, title_value = cells_in_row[0]
        if label_col > column:
            continue
        if not isinstance(title_value, str):
            continue
        text = title_value.strip().strip("'\"")
        if not text or text.startswith("="):
            continue
        if len(text) > 16 or re.search(r"20\d{2}[AaEe]?", text) or re.fullmatch(r"20\d{2}", text):
            continue
        normalized = text.replace(" ", "").lower()
        if any(token in normalized for token in _metric_tokens):
            continue
        if any(token in normalized for token in _header_tokens) and len(normalized) <= 6:
            continue
        # 同行右侧必须没有数值 — 数据行的标签不是区块标题
        has_value_right = any(isinstance(value, (int, float)) and not isinstance(value, bool)
                              for _, value in cells_in_row[1:])
        if has_value_right:
            continue
        # 目标行必须在标题行下方 (标题不能与目标同行)
        if candidate_row >= row:
            continue
        return text
    return ""


def _nearest_header(column_index, row, column):
    """Collect the full header path above this cell (multi-row headers).

    Spreadsheet summary matrices frequently stack a year row above an entity/
    scenario row (e.g. 2022A / 无锡|东台|合并). Returning only the nearest
    single header collapses distinct columns into one period key and fabricates
    conflicts. Headers are read from the top of the column (sheet header zone,
    rows 1..10); stray text near the data row (labels, formulas, notes) must
    never shadow the real column header.
    """
    entries = sorted((candidate_row, value) for candidate_row, value
                     in column_index.get(column, ()) if candidate_row < row)
    header_rows = [entry for entry in entries if entry[0] <= 10]
    if not header_rows:
        header_rows = entries[:6]
    path, fallback = [], ""
    for candidate_row, value in reversed(header_rows):
        if isinstance(value, (datetime, date)):
            text = value.isoformat()
        elif isinstance(value, str) and not value.startswith("="):
            text = value.strip().strip("'\"")
        else:
            continue
        if not text:
            continue
        if re.search(r"\b20\d{2}\s*[AaEe]?\b", text):
            return "·".join((text, *path)) if path else text
        if len(path) < 2:
            path.insert(0, text)
        fallback = fallback or text
    return "·".join(path) if path else fallback


def _nearest_column_label(column_index, row, column):
    for candidate_row, value in reversed(column_index.get(column, ())):
        if candidate_row >= row:
            continue
        if row - candidate_row > 10:
            break
        if isinstance(value, str) and value.strip() and not value.startswith("="):
            text = value.strip()
            if not re.fullmatch(r"20\d{2}[AaEe]?", text):
                return text
    return ""


def _parent_metric(row_index, row, column):
    for candidate_row in range(row - 1, max(0, row - 4), -1):
        for candidate_col, value in reversed(row_index.get(candidate_row, ())):
            if candidate_col >= column:
                continue
            if isinstance(value, str) and value.strip() and not value.startswith("="):
                return value.strip()
    return ""


def _sheet_unit(cells):
    for cell in cells:
        if cell.row > 10:
            break
        raw_text = str(cell.value or "")
        text = raw_text.casefold()
        if "million" in text and ("chinese yuan" in text or "rmb" in text):
            return "百万元人民币"
        for unit in ("百万元", "万元", "亿元", "人民币元"):
            if unit in raw_text and ("单位" in raw_text or "币种" in raw_text):
                return unit
    return None


def _cell_unit(cell, label, sheet_unit):
    number_format = str(cell.number_format).casefold().replace("\\", "")
    if "%" in number_format:
        return "%"
    if "x" in number_format:
        return "x"
    normalized_label = str(label).replace(" ", "")
    if "元/股" in normalized_label or "元／股" in normalized_label:
        return "元/股"
    return sheet_unit


def _formula_text(value):
    if isinstance(value, str):
        return value
    kind = getattr(value, "t", type(value).__name__)
    attributes = []
    for name in ("ref", "text", "r1", "r2", "dt2D", "dtr"):
        item = getattr(value, name, None)
        if item not in (None, False, ""):
            attributes.append(f"{name}={item}")
    return f"<{kind}:{';'.join(attributes)}>"


def _instantiated_cells(sheet):
    # openpyxl has no public sparse-cell iterator. Reading the already loaded cell
    # store avoids catastrophic max_row x max_column scans on formatted full sheets.
    return sorted((cell for cell in sheet._cells.values() if cell.value is not None),
                  key=lambda cell: (cell.row, cell.column))


def _sparse_indexes(cells):
    rows, columns = defaultdict(list), defaultdict(list)
    for cell in cells:
        rows[cell.row].append((cell.column, cell.value))
        columns[cell.column].append((cell.row, cell.value))
    return rows, columns


def _cell_value(sheet, row, column):
    cell = sheet._cells.get((row, column))
    return None if cell is None else cell.value


def _key_value(text):
    for separator in (":", "："):
        if separator in text:
            label, value = text.split(separator, 1)
            return label.strip(), _scalar(value.strip())
    return text, text


def _scalar(value):
    if isinstance(value, (datetime, date)):
        return value.isoformat()
    if isinstance(value, (int, float, bool)) or value is None:
        return value
    if not isinstance(value, str):
        return str(value)
    text = value.strip().replace(",", "")
    try:
        return float(text) if "." in text else int(text)
    except ValueError:
        return value.strip()


def _item(path, source_type, location, label, value, raw="", formula=None, cached=None, context=(), unit=None,
          coerce_value=True):
    normalized_value = _scalar(value) if coerce_value else value
    return EvidenceItem(str(path), source_type, location, str(label).strip(), normalized_value, unit=unit,
                        formula=formula, cached_value=_scalar(cached), raw_text=str(raw), context=tuple(context))


_LOADERS = {".csv": _csv, ".json": _json, ".docx": _docx, ".xlsx": _xlsx, ".pdf": _pdf}
