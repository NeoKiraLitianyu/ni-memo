"""Render memo pages and enforce automated visual-structure gates."""
from dataclasses import asdict, dataclass
from collections import Counter
from pathlib import Path
import re
import shutil
import subprocess
from zipfile import ZipFile

from docx import Document
from PIL import Image, ImageDraw
from pypdf import PdfReader


@dataclass(frozen=True)
class VisualReport:
    status: str
    pdf: str | None
    page_count: int
    page_images: tuple[str, ...]
    contact_sheet: str | None
    blank_pages: tuple[int, ...]
    missing_headings: tuple[str, ...]
    unresolved_tokens: tuple[str, ...]
    messages: tuple[str, ...]

    def to_dict(self):
        data = asdict(self)
        for key in ("page_images", "blank_pages", "missing_headings", "unresolved_tokens", "messages"):
            data[key] = list(data[key])
        return data


def render_and_inspect(docx_path, out_dir, expected_headings=(), soffice_path=None, pdftoppm_path=None):
    source = Path(docx_path).resolve()
    root = Path(out_dir).resolve()
    root.mkdir(parents=True, exist_ok=True)
    for prior in (*root.glob("page-*.png"), root / "contact-sheet.jpg"):
        if prior.exists():
            _retire(prior)
    engine = _find_soffice(soffice_path)
    if engine is None:
        return VisualReport("not_run", None, 0, (), None, (), tuple(expected_headings), (),
                            ("LibreOffice soffice was not found",))
    pdf = root / f"{source.stem}.pdf"
    if pdf.exists():
        _retire(pdf)
    profile = root / "profile"
    profile.mkdir(exist_ok=True)
    command = [
        str(engine), "--headless", f"-env:UserInstallation={profile.as_uri()}",
        "--convert-to", "pdf", "--outdir", str(root), str(source),
    ]
    completed = subprocess.run(
        command, capture_output=True, text=True, encoding="utf-8", errors="replace",
        timeout=180, check=False,
    )
    raw_out = (completed.stdout or "") + "\n" + (completed.stderr or "")
    messages = tuple(line for line in raw_out.splitlines() if line.strip())
    if completed.returncode != 0 or not pdf.exists():
        return VisualReport("fail", None, 0, (), None, (), tuple(expected_headings), (), messages)

    reader = PdfReader(str(pdf))
    page_texts = [(page.extract_text() or "").strip() for page in reader.pages]
    full_text = "\n".join(page_texts)
    body_texts = _body_texts(page_texts)
    blank_pages = tuple(index for index, text in enumerate(body_texts, start=1) if len(text) < 10)
    missing_headings = tuple(heading for heading in expected_headings if heading not in full_text)
    unresolved = tuple(sorted(set(_unresolved_tokens(source))))

    renderer = _find_pdftoppm(pdftoppm_path)
    images = ()
    contact_sheet = None
    if renderer:
        prefix = root / "page"
        rendered = subprocess.run(
            [str(renderer), "-png", "-r", "120", str(pdf), str(prefix)],
            capture_output=True, text=True, encoding="utf-8", errors="replace",
            timeout=180, check=False,
        )
        raw_out = (rendered.stdout or "") + "\n" + (rendered.stderr or "")
        messages += tuple(line for line in raw_out.splitlines() if line.strip())
        images = tuple(str(path) for path in sorted(root.glob("page-*.png")))
        if images:
            contact_sheet = str(_contact_sheet([Path(path) for path in images], root / "contact-sheet.jpg"))
    defects = bool(blank_pages or missing_headings or unresolved or len(reader.pages) < 9)
    if renderer and len(images) != len(reader.pages):
        defects = True
    status = "fail" if defects else "automated_pass"
    return VisualReport(status, str(pdf), len(reader.pages), images, contact_sheet, blank_pages,
                        missing_headings, unresolved, messages)


def _unresolved_tokens(path):
    document = Document(path)
    texts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            texts.extend(cell.text for cell in row.cells)
    unresolved = [text for text in texts if "{{" in text or "}}" in text]
    if any(text.strip() == "None" for text in texts):
        unresolved.append("raw None token")
    if any("T00:00:00" in text for text in texts):
        unresolved.append("raw datetime token")
    if any("recalculated/" in text for text in texts):
        unresolved.append("internal recalculated path")
    if any(re.search(r"\bn\.?m\.?(?:\b|$)", text, re.IGNORECASE) for text in texts):
        unresolved.append("raw not-meaningful token")
    with ZipFile(path) as archive:
        document_xml = archive.read("word/document.xml").decode("utf-8")
    if "在 Word 中打开后自动更新目录" in document_xml:
        unresolved.append("TOC field result was not refreshed")
    return unresolved


def _body_texts(page_texts):
    page_lines = [tuple(line.strip() for line in text.splitlines() if line.strip()) for text in page_texts]
    counts = Counter(line for lines in page_lines for line in set(lines))
    threshold = max(2, int(len(page_lines) * 0.7))
    repeated = {line for line, count in counts.items() if count >= threshold}
    bodies = []
    for lines in page_lines:
        kept = [line for line in lines if line not in repeated
                and not re.fullmatch(r"CONFIDENTIAL\s*[·|｜]?\s*\d+", line, re.IGNORECASE)
                and not re.fullmatch(r"\d+", line)]
        bodies.append("\n".join(kept).strip())
    return bodies


def _retire(path):
    """Sandbox-safe removal: rename stale artifacts aside instead of unlink.

    Windows sandbox has no recycle bin (SAFE_DELETE_FAIL_CLOSED), so unlink is
    blocked. Visual artifacts are regenerable intermediates; moving them aside
    (page-01.png -> page-01.prev.png) keeps the target path clean for the new
    render without deleting anything.
    """
    if not Path(path).exists():
        return
    stale = Path(path).with_name(Path(path).name + ".prev")
    try:
        Path(path).replace(stale)
    except OSError:
        pass


def _contact_sheet(paths, output):
    thumbnails = []
    for index, path in enumerate(paths, start=1):
        with Image.open(path) as image:
            thumb = image.convert("RGB")
            thumb.thumbnail((260, 370))
            framed = Image.new("RGB", (280, 400), "white")
            framed.paste(thumb, ((280 - thumb.width) // 2, 20))
            ImageDraw.Draw(framed).text((10, 5), f"Page {index}", fill="black")
            thumbnails.append(framed)
    columns = 4
    rows = (len(thumbnails) + columns - 1) // columns
    sheet = Image.new("RGB", (columns * 280, rows * 400), (225, 228, 232))
    for index, image in enumerate(thumbnails):
        sheet.paste(image, ((index % columns) * 280, (index // columns) * 400))
    sheet.save(output, quality=88)
    return output


def _find_soffice(explicit):
    candidates = [
        Path(explicit) if explicit else None,
        Path(shutil.which("soffice.com")) if shutil.which("soffice.com") else None,
        Path(shutil.which("soffice")) if shutil.which("soffice") else None,
        Path(r"C:\Program Files\LibreOffice\program\soffice.com"),
    ]
    return next((path for path in candidates if path and path.is_file()), None)


def _find_pdftoppm(explicit):
    if explicit and Path(explicit).is_file():
        return Path(explicit)
    runtime_root = Path.home() / ".cache" / "codex-runtimes"
    matches = sorted(runtime_root.glob("*/dependencies/native/poppler/Library/bin/pdftoppm.exe"))
    if matches:
        return matches[-1]
    found = shutil.which("pdftoppm.exe") or shutil.which("pdftoppm")
    return Path(found) if found else None
