"""Render stable fact tokens without importing project binders."""
from dataclasses import dataclass
from datetime import datetime
from io import BytesIO
from pathlib import Path
import re
import shutil
import zipfile
from urllib.parse import urlparse

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from openpyxl.utils.cell import column_index_from_string, get_column_letter

from ni_memo.chart import render_series_chart_png
from ni_memo.quality import evaluate_completion

# ── 海古德 memo 版式规格 (docs/haigude_layout_spec.md) ─────────────────────
# 正文: A4 纵向 11906 × 16838 twips; 边距: 上下 1440 / 左右 1800 (twips)
PAGE_W = Inches(11906 / 1440)      # ≈ 8.27 in (A4 短边)
PAGE_H = Inches(16838 / 1440)      # ≈ 11.69 in (A4 长边)
MARGIN_TOP_BOTTOM = Inches(1440 / 1440)
MARGIN_LEFT_RIGHT = Inches(1800 / 1440)
# 附录可比表可横向放宽: 16838 × 11906, 边距 1797/1440
LANDSCAPE_W = Inches(16838 / 1440)
LANDSCAPE_H = Inches(11906 / 1440)
LANDSCAPE_TB = Inches(1797 / 1440)
LANDSCAPE_LR = Inches(1440 / 1440)
# 字体: Latin = Palatino Linotype, 中文 = 仿宋 (Normal sz=22 → 11pt)
LATIN_FONT = "Palatino Linotype"
EAST_ASIA_FONT = "仿宋"
# heading1: sz=28 → 14pt + 下边框 single sz=6; heading2: sz=24 → 12pt 加粗
TOKEN = re.compile(r"\{\{([a-zA-Z0-9_.-]+)\}\}")


@dataclass(frozen=True)
class RenderReport:
    snapshot_id: str
    replacements: int
    output_path: str


def resolve_token(text, snapshot):
    def replace(match):
        key = match.group(1)
        if key not in snapshot.selected:
            raise KeyError(f"missing fact token: {key}")
        fact = snapshot.selected[key]
        return f"{_display_value(fact.value)}{fact.unit or ''}"
    return TOKEN.sub(replace, text)


def render_docx(template_path, output_path, snapshot):
    doc = Document(template_path)
    replacements = 0
    paragraphs = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs.extend(cell.paragraphs)
    for paragraph in paragraphs:
        for run in paragraph.runs:
            count = len(TOKEN.findall(run.text))
            if count:
                run.text = resolve_token(run.text, snapshot)
                replacements += count
    doc.save(output_path)
    return RenderReport(snapshot.snapshot_id, replacements, str(output_path))


def render_default_docx(output_path, snapshot, project_name, grade):
    """Create a conservative memo when the caller supplies data but no template."""
    doc = Document()
    section = doc.sections[0]
    section.page_width, section.page_height = Inches(8.5), Inches(11)
    section.top_margin = section.bottom_margin = Inches(0.8)
    section.left_margin = section.right_margin = Inches(0.85)
    _style_document(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(title, f"{project_name} Investment Memo", 22, True, RGBColor(31, 77, 120))
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(subtitle, f"Source-traceable model snapshot | {grade}", 10, False, RGBColor(90, 98, 108))

    doc.add_heading("1. Executive summary", level=1)
    doc.add_paragraph(
        "This memo was generated from the supplied financial model. It records model facts and exact source lineage; "
        "source-only values are not represented as independently verified due-diligence conclusions."
    )
    doc.add_heading("2. Transaction and valuation", level=1)
    _fact_table(doc, snapshot, ("valuation.", "investment."))
    doc.add_heading("3. Financial snapshot", level=1)
    _fact_table(doc, snapshot, ("financial.", "balance."))
    doc.add_heading("4. Evidence status and limitations", level=1)
    doc.add_paragraph(f"Snapshot ID: {snapshot.snapshot_id}")
    doc.add_paragraph(f"Acceptance grade: {grade}")
    doc.add_paragraph(
        "All figures retain workbook sheet/cell lineage. Formula-derived figures use the workbook's cached result "
        "when available and retain the original formula in evidence. Units remain 'model unit' when the workbook "
        "does not provide an unambiguous currency scale."
    )
    doc.save(output_path)
    _pin_core_modified(output_path)
    return RenderReport(snapshot.snapshot_id, len(snapshot.selected), str(output_path))


def render_reference_docx(reference_path, output_path, snapshot_id):
    """Preserve an explicitly supplied historical DOCX without rewriting its XML."""
    source = Path(reference_path)
    if not source.is_file() or source.suffix.lower() != ".docx":
        raise ValueError("reference_docx must be an existing DOCX file")
    target = Path(output_path)
    target.parent.mkdir(parents=True, exist_ok=True)
    shutil.copyfile(source, target)
    return RenderReport(snapshot_id, 0, str(target))


def render_standard_memo(output_path, schema, snapshot, quality, formula, project_name,
                         narratives=None, workbook_profiles=(), analysis=None, completion=None,
                         project_as_of=None):
    """Render the Haigude-layout investment memo (A4 portrait, 仿宋 11pt).

    Layout parity with 臻芯基金-海古德项目memo-20240927FV.docx:
      * A4 portrait, margins 1440/1800 twips (appendix tables can be landscape)
      * Normal: Palatino Linotype + 仿宋 11pt, justified
      * Cover: 致/日期/有关 three-line block with bottom border
      * TOC: updateable Word field covering Heading 1-2
      * Chapters: roman-numeral numbered, bold 14pt, bottom border
      * Sections: 中文数字 numbered (一、二、…), bold 12pt
      * Sub-headings: a9 List Paragraph (2-char indent)
      * Tables: Table Grid full borders 0.5pt, header bold, no shading
      * Footer: page number only; header empty
      * Raw workbook regions remain in the audit work directory, never the DOCX
    """
    doc = Document()
    section = doc.sections[0]
    _setup_page(section)
    _style_standard_document(doc)
    doc.core_properties.title = f"{project_name} 投资备忘录"

    # ── 封面头部 (致/日期/有关) ──
    cover_to = _memo_cover_para(doc)
    _add_run(cover_to, "致：", 14, True, None, east="仿宋")
    _add_run(cover_to, "投资决策委员会", 14, True, None, east="仿宋")
    cover_date = _memo_cover_para(doc)
    _add_run(cover_date, f"日期：{datetime.now().year}年{datetime.now().month}月", 12, False, None, east="仿宋")
    if project_as_of:
        cover_as_of = _memo_cover_para(doc)
        _add_run(cover_as_of, f"项目资料截止日：{project_as_of}", 10, False, None, east="仿宋")
    cover = _memo_cover_para(doc)
    _add_run(cover, "有关：", 12, False, None, east="仿宋")
    _add_run(cover, project_name, 12, False, None, east="仿宋")
    _add_run(cover, "投资项目", 12, False, None, east="仿宋")
    _cover_bottom_border(cover)
    doc.add_paragraph()

    # ── 目录 (真实 Word TOC 字段; 打开文档时由 Word 更新页码) ──
    _render_toc(doc, schema)

    # ── 章节正文 (typed content blocks: facts + narrative + workbook regions) ──
    completion = completion or evaluate_completion(schema, snapshot)
    for index, section_def in enumerate(schema.sections, start=1):
        if index > 1 and section_def.new_word_section:
            _configure_formal_section(
                doc.add_section(WD_SECTION.NEW_PAGE), section_def.orientation,
            )
        chapter = doc.add_paragraph(style="Heading 1")
        chapter.paragraph_format.page_break_before = True
        _add_run(chapter, f"{_roman(index)} {section_def.title}", 14, True, None, east="仿宋")
        _heading1_border(chapter)
        fields = [field for field in schema.fields if field.section == section_def.key]
        blocks = [block for block in schema.blocks if block.section == section_def.key]
        for sub_index, block in enumerate(blocks, start=1):
            _render_block(doc, block, sub_index, fields, snapshot, schema, narratives,
                          analysis_items=analysis.get(block.key) if analysis else (),
                          completion=completion, quality=quality, formula=formula)

    _enforce_table_controls(doc)
    _enable_field_updates(doc)
    doc.save(output_path)
    _pin_core_modified(output_path)
    return RenderReport(snapshot.snapshot_id, len(snapshot.selected), str(output_path))


def _pin_core_modified(path):
    """Make the DOCX byte-deterministic: python-docx writes the current time into
    `docProps/core.xml` dcterms:modified and into every zip entry's DOS timestamp on
    save, so two renders of the same snapshot differ in bytes despite identical
    content. Rewrite the modified element to a fixed value and pin every zip entry
    timestamp to a fixed epoch so the package is byte-for-byte reproducible."""
    p = Path(path)
    fixed = b"2024-09-27T00:00:00Z"  # 参照 memo 基准日期, 与正文 as_of 一致
    # DOS date/time: 2024-09-27 = (year-1980)<<9 | month<<5 | day = 44<<9|9<<5|27
    dos_datetime = (44 << 9) | (9 << 5) | 27  # 2024-09-27, time 00:00:00
    tmp = p.with_suffix(".deterministic.tmp")
    with zipfile.ZipFile(p, "r") as zin, zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename == "docProps/core.xml":
                text = data.decode("utf-8")
                text = re.sub(
                    r"<dcterms:modified[^>]*>.*?</dcterms:modified>",
                    f"<dcterms:modified xsi:type=\"dcterms:W3CDTF\">{fixed.decode()}</dcterms:modified>",
                    text, flags=re.S,
                )
                data = text.encode("utf-8")
            item.date_time = (2024, 9, 27, 0, 0, 0)
            zout.writestr(item, data)
    tmp.replace(p)


def _setup_page(section):
    _configure_formal_section(section, "portrait")
    section.header_distance = Inches(851 / 1440)
    section.footer_distance = Inches(992 / 1440)
    # 页眉留空 (海古德 header 为空), 页脚仅居中页码
    _footer_page_number(section.footer.paragraphs[0])


def _configure_formal_section(section, orientation):
    if orientation == "landscape":
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = LANDSCAPE_W
        section.page_height = LANDSCAPE_H
        section.top_margin = section.bottom_margin = LANDSCAPE_TB
        section.left_margin = section.right_margin = LANDSCAPE_LR
    else:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = PAGE_W
        section.page_height = PAGE_H
        section.top_margin = section.bottom_margin = MARGIN_TOP_BOTTOM
        section.left_margin = section.right_margin = MARGIN_LEFT_RIGHT
    section.header_distance = Inches(851 / 1440)
    section.footer_distance = Inches(992 / 1440)


def _memo_cover_para(doc):
    para = doc.add_paragraph()
    para.paragraph_format.line_spacing = 1.15
    para.paragraph_format.space_after = Pt(3)
    return para


def _cover_bottom_border(paragraph):
    ppr = paragraph._p.get_or_add_pPr()
    ppr.append(_borders_el("single", "6", "2"))


def _heading1_border(paragraph):
    ppr = paragraph._p.get_or_add_pPr()
    ppr.append(_borders_el("single", "6", "1"))


def _borders_el(val, sz, space):
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), val)
    bottom.set(qn("w:sz"), sz)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), "auto")
    borders.append(bottom)
    return borders


def _footer_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE \\* MERGEFORMAT "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instruction, separate, text, end))


def _render_toc(doc, schema):
    """Insert a real TOC field; never guess page numbers in generated text."""
    heading = doc.add_paragraph(style="TOC Heading")
    _add_run(heading, "目录", 16, True, RGBColor(0x2F, 0x54, 0x96), east="仿宋")
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin.set(qn("w:dirty"), "true")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = ' TOC \\o "1-2" \\h \\z \\u '
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "在 Word 中打开后自动更新目录"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instruction, separate, placeholder, end))
    doc.add_page_break()


def _enable_field_updates(doc):
    settings = doc.settings._element
    existing = settings.find(qn("w:updateFields"))
    update = existing if existing is not None else OxmlElement("w:updateFields")
    update.set(qn("w:val"), "true")
    if existing is None:
        settings.append(update)


def _render_block(doc, block, sub_index, fields, snapshot, schema, narratives,
                  analysis_items=(), completion=None, quality=None, formula=None):
    """Render one schema block as a 节 (一、二、… heading + content)."""
    heading = doc.add_paragraph(style="Heading 2")
    _add_run(heading, f"{_chinese_number(sub_index)} {block.title}", 12, True, None, east="仿宋")
    heading.paragraph_format.space_before = Pt(10)
    selected = _facts_for_block(block, fields, snapshot)
    # narrative_key 未显式声明时 fallback 到 block.key, 让 financial.* 等
    # 未声明 narrative_key 的块也能消费 narrative.json 侧车素材 (2.1.0 语义)。
    narrative_key = block.narrative_key or block.key
    explicit_narratives = narratives.get(narrative_key) if narratives is not None else ()
    block_narratives = tuple(explicit_narratives) + tuple(analysis_items or ())
    if block.key == "appendix.completion_review":
        _completion_review_block(doc, completion, quality, formula)
        return
    if block_narratives:
        _narrative_items_block(doc, block_narratives)
    # 叙述素材与结构化内容的关系:
    # - narrative 块: 素材替代事实叙述 (无素材时才用事实生成叙述)
    # - financial_table/returns_table/sensitivity_table: 素材 + 财务表格**并存**
    #   — 对齐参照版 20240927FV 的"财务情况/财务预测"章节"叙述段落 + 数据表格"
    #   双层结构
    # - 其他块: 素材优先 (有素材不渲染事实内容, 维持既有行为)
    if block.block_type == "narrative":
        if not block_narratives:
            _narrative_block(doc, selected, schema, snapshot, block.title)
    elif block.block_type in {"financial_table", "returns_table", "sensitivity_table"} and selected:
        _series_fact_table(doc, selected, schema, snapshot)
    elif block.block_type in {"bullet_list", "risk_cards", "product_profile"} and selected:
        _bullet_block(doc, selected, schema, snapshot, risk=block.block_type == "risk_cards")
    elif block.key == "appendix.sources" and selected:
        _source_index_block(doc, selected, schema, snapshot)
    elif selected:
        _standard_fact_table(doc, selected, schema, snapshot)
    elif not block_narratives:
        _missing_block(doc, block.title)
    if selected and block.block_type in {"chart", "financial_table"}:
        _semantic_series_chart(doc, block, selected, schema, snapshot)


def _roman(index):
    values = {1: "I", 2: "II", 3: "III", 4: "IV", 5: "V", 6: "VI",
              7: "VII", 8: "VIII", 9: "IX", 10: "X"}
    return values.get(index, str(index))


def _semantic_series_chart(doc, block, facts, schema, snapshot):
    grouped = {}
    for key, fact in facts:
        period = _period(key)
        if not period or not isinstance(fact.value, (int, float)) or isinstance(fact.value, bool):
            continue
        field = _field_for_key(schema, key)
        grouped.setdefault(field.unit_family, {}).setdefault(field.key, {})[period] = float(fact.value)
    candidates = []
    for unit_family, fields in grouped.items():
        for field_key, values in fields.items():
            if len(values) >= 2:
                candidates.append((len(values), unit_family, field_key, values))
    if not candidates:
        return
    _, unit_family, primary_key, primary_values = max(candidates, key=lambda item: item[0])
    periods = tuple(sorted(primary_values, key=_period_sort_key))
    series = []
    for field_key, values in grouped[unit_family].items():
        if all(period in values for period in periods):
            field = next(item for item in schema.fields if item.key == field_key)
            series.append((field.label, tuple(values[period] for period in periods)))
    chart = render_series_chart_png(periods, tuple(series))
    if not chart:
        return
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.keep_with_next = True
    _add_run(caption, f"图：{block.title}（基于冻结事实）", 9, True, None, east="仿宋")
    picture = doc.add_paragraph()
    picture.alignment = WD_ALIGN_PARAGRAPH.CENTER
    picture.add_run().add_picture(BytesIO(chart), width=Inches(5.9))
    sources = tuple(dict.fromkeys(
        _fact_sources(key, fact, snapshot)
        for key, fact in facts if _period(key) in periods
    ))
    source = doc.add_paragraph()
    source.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(
        source, f"图表依据：{'；'.join(sources)}", 7.5, False,
        RGBColor(105, 113, 122),
    )


def _repeat_header(row):
    properties = row._tr.get_or_add_trPr()
    if properties.find(qn("w:tblHeader")) is not None:
        return
    marker = OxmlElement("w:tblHeader")
    marker.set(qn("w:val"), "true")
    properties.append(marker)


def _keep_row_together(row):
    properties = row._tr.get_or_add_trPr()
    if properties.find(qn("w:cantSplit")) is not None:
        return
    properties.append(OxmlElement("w:cantSplit"))


def _enforce_table_controls(doc):
    for table in doc.tables:
        if not table.rows:
            continue
        _repeat_header(table.rows[0])
        for row in table.rows[1:]:
            _keep_row_together(row)


def _shade_cell(cell, fill):
    properties = cell._tc.get_or_add_tcPr()
    shade = properties.find(qn("w:shd"))
    if shade is None:
        shade = OxmlElement("w:shd")
        properties.append(shade)
    shade.set(qn("w:fill"), fill)


def _chinese_number(index):
    values = {1: "一", 2: "二", 3: "三", 4: "四", 5: "五", 6: "六",
              7: "七", 8: "八", 9: "九", 10: "十"}
    return values.get(index, str(index))


def _style_document(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    for name, size in (("Heading 1", 15), ("Heading 2", 12)):
        style = doc.styles[name]
        style.font.name = "Aptos Display"
        style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor(46, 116, 181)


def _style_standard_document(doc):
    """Haigude layout: Palatino Linotype + 仿宋 11pt justified; no colored headings."""
    normal = doc.styles["Normal"]
    normal.font.name = LATIN_FONT
    normal._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
    normal.font.size = Pt(11)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(4)
    # The formal reference inherits Latin/CJK fonts from Normal for both headings.
    # Runs carry the visible font and weight, while the styles keep the reference
    # inheritance contract (Heading 1 weight inherited; Heading 2 explicitly bold).
    for name, size in (("Heading 1", 14), ("Heading 2", 12)):
        style = doc.styles[name]
        style.font.name = None
        fonts = style._element.get_or_add_rPr().rFonts
        if fonts is not None:
            fonts.attrib.pop(qn("w:eastAsia"), None)
        style.font.size = Pt(size)
        style.font.bold = True if name == "Heading 2" else None
        style.font.color.rgb = RGBColor(0, 0, 0)
    doc.styles["Heading 1"].paragraph_format.space_before = Pt(6)
    doc.styles["Heading 1"].paragraph_format.space_after = Pt(6)
    doc.styles["Heading 2"].paragraph_format.space_before = Pt(8)
    doc.styles["Heading 2"].paragraph_format.space_after = Pt(4)
    # TOC 样式 (TOC Heading / TOC1 / TOC2) — python-docx 默认模板无 TOC 样式, 动态创建
    toc_heading = _ensure_style(doc, "TOC Heading")
    toc_heading.font.name = LATIN_FONT
    toc_heading._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
    toc_heading.font.size = Pt(16)
    toc_heading.font.bold = True
    for name in ("TOC 1", "TOC 2"):
        style = _ensure_style(doc, name)
        style.font.name = LATIN_FONT
        style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
        style.font.size = Pt(11)
    # Table Grid 全边框
    try:
        grid = doc.styles["Table Grid"]
        grid.font.name = LATIN_FONT
        grid._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
    except KeyError:
        pass


def _ensure_style(doc, name):
    """Return a named paragraph style, creating a Normal-derived one if absent."""
    try:
        return doc.styles[name]
    except KeyError:
        from docx.enum.style import WD_STYLE_TYPE
        return doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)


def _add_header_footer(section, project_name):
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _add_run(header, f"臻芯基金｜{project_name}｜投资备忘录", 8, False, RGBColor(105, 113, 122))
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(footer, "CONFIDENTIAL  ·  ", 7.5, False, RGBColor(120, 126, 132))
    run = footer.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instruction, end))


def _status_table(doc, grade, formula_status, snapshot_id):
    table = doc.add_table(rows=3, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    rows = (("文档等级", grade), ("公式门禁", formula_status), ("事实快照", snapshot_id))
    for row, values in zip(table.rows, rows):
        for cell, value in zip(row.cells, values):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            _add_run(cell.paragraphs[0], str(value), 8.5, cell is row.cells[0])


def _facts_for_block(block, fields, snapshot):
    prefixes = block.field_prefixes
    if block.key == "appendix.sources":
        return [(key, snapshot.selected[key]) for key in sorted(snapshot.selected)]
    field_keys = {field.key for field in fields}
    facts = [(key, snapshot.selected[key]) for key in sorted(snapshot.selected)
             if any(key == prefix or key.startswith(prefix) for prefix in prefixes)
             and any(key == field_key or key.startswith(field_key + ".") for field_key in field_keys)]
    if block.key == "financial.historical":
        return [(key, fact) for key, fact in facts if _period(key).endswith("A")]
    if block.key == "financial.forecast":
        return [(key, fact) for key, fact in facts if _period(key).endswith("E")]
    return facts


def _narrative_block(doc, facts, schema, snapshot, title="本节"):
    if not facts:
        _missing_block(doc, title)
        return
    summary = "；".join(
        f"{_field_for_key(schema, key).label}：{_fact_value(key, fact, schema, snapshot)}"
        for key, fact in facts[:6]
    )
    doc.add_paragraph(f"根据已提供资料，本节可确认的核心事实包括：{summary}。其余判断需结合来源状态与待尽调事项阅读。")


def _missing_block(doc, title):
    paragraph = doc.add_paragraph()
    _add_run(paragraph, f"资料未提供：尚无可唯一绑定到“{title}”的有效证据，未作推断或编造。", 9.5,
             False, RGBColor(156, 92, 38))


def _narrative_items_block(doc, items):
    """Render user-supplied investment narrative (highlight/risk/industry/company).

    Mirrors the Haigude memo style: each item is a bold title followed by a
    substantive argument paragraph. Narrative comes from the caller's JSON sidecar,
    never fabricated by the engine.
    """
    for item in items:
        if item.title:
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.space_before = Pt(8)
            paragraph.paragraph_format.space_after = Pt(3)
            _add_run(paragraph, item.title, 10.5, True, RGBColor(21, 40, 64))
        body = doc.add_paragraph()
        body.paragraph_format.space_after = Pt(7)
        _add_run(body, item.body, 9.5, False, RGBColor(40, 44, 50))
        sources = tuple(getattr(item, "sources", ()))
        if sources:
            source = doc.add_paragraph()
            source.paragraph_format.left_indent = Inches(0.2)
            source.paragraph_format.space_after = Pt(6)
            label = "模型依据" if getattr(item, "inference", None) else "叙事依据"
            _add_run(
                source, f"{label}：{_display_claim_sources(sources)}", 7.5, False,
                RGBColor(105, 113, 122),
            )


def _display_claim_sources(sources):
    displayed = []
    for source in sources:
        uri, separator, location = str(source).rpartition("#")
        value = f"{_claim_source_name(uri)}#{location}" if separator else _claim_source_name(source)
        displayed.append(value)
    return "；".join(dict.fromkeys(displayed))


def _claim_source_name(uri):
    parsed = urlparse(str(uri))
    if parsed.scheme in {"http", "https"}:
        return Path(parsed.path).name or parsed.netloc
    return Path(str(uri)).name


def _completion_review_block(doc, completion, quality, formula):
    verified_rate = float(getattr(quality, "metrics", {}).get("verified_rate", 0.0))
    formula_status = getattr(formula, "status", "not_applicable")
    summary = doc.add_table(rows=6, cols=2)
    summary.alignment = WD_TABLE_ALIGNMENT.CENTER
    summary.style = "Table Grid"
    rows = (
        ("核心必填字段完整度", f"{completion.required_rate:.1f}% / 100%（{completion.required_completed}/{completion.required_total}）"),
        ("必需槽位完整度", f"{completion.required_slot_rate:.1f}% / 100%（{completion.required_slots_completed}/{completion.required_slots_total}）"),
        ("全字段覆盖率", f"{completion.all_field_rate:.1f}% / 100%（{completion.all_completed}/{completion.all_total}）"),
        ("全槽位覆盖率", f"{completion.all_slot_rate:.1f}% / 100%（{completion.all_slots_completed}/{completion.all_slots_total}）"),
        ("独立验证率", f"{verified_rate:.1f}% / 100%"),
        ("公式门禁", formula_status),
    )
    for row, values in zip(summary.rows, rows):
        for index, (cell, value) in enumerate(zip(row.cells, values)):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            _add_run(cell.paragraphs[0], value, 8.5, index == 0)

    heading = doc.add_paragraph()
    heading.paragraph_format.space_before = Pt(8)
    _add_run(heading, "逐字段完成状态", 9.5, True, RGBColor(21, 40, 64))
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    _repeat_header(table.rows[0])
    for cell, value in zip(table.rows[0].cells, ("资料项", "要求", "完成状态", "证据状态")):
        _shade_cell(cell, "1F4E78")
        _add_run(cell.paragraphs[0], value, 8, True, RGBColor(255, 255, 255))
    status_labels = {
        "complete": "完成", "partial": "部分完成", "conflict": "冲突",
        "later_update": "后续公开更新", "missing": "缺失",
    }
    evidence_labels = {
        "verified": "已核验", "source_only": "仅来源可追溯", "retained_original": "历史原文保留",
        "mixed": "混合证据", "unresolved": "待解决", "none": "无证据",
        "editorial_sourced": "有来源的编辑叙事", "source_index": "来源索引",
        "field_states": "字段状态汇总",
    }
    for item in completion.items:
        row = table.add_row()
        _keep_row_together(row)
        cells = row.cells
        values = (
            item.label,
            "必填" if item.requirement != "optional" else "可选",
            status_labels[item.status],
            evidence_labels[item.evidence_status],
        )
        for cell, value in zip(cells, values):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            _add_run(cell.paragraphs[0], value, 7.5)

    slot_heading = doc.add_paragraph()
    slot_heading.paragraph_format.space_before = Pt(8)
    _add_run(slot_heading, "逐槽位完成状态", 9.5, True, RGBColor(21, 40, 64))
    slot_table = doc.add_table(rows=1, cols=4)
    slot_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    slot_table.style = "Table Grid"
    _repeat_header(slot_table.rows[0])
    for cell, value in zip(slot_table.rows[0].cells, ("正式模板槽位", "要求", "完成状态", "依据")):
        _shade_cell(cell, "1F4E78")
        _add_run(cell.paragraphs[0], value, 8, True, RGBColor(255, 255, 255))
    for item in completion.slots:
        row = slot_table.add_row()
        _keep_row_together(row)
        values = (
            item.label,
            "必需" if item.requirement == "required" else "可选",
            status_labels[item.status],
            evidence_labels.get(item.evidence_status, item.evidence_status),
        )
        for cell, value in zip(row.cells, values):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            _add_run(cell.paragraphs[0], value, 7.5)

    missing_heading = doc.add_paragraph()
    missing_heading.paragraph_format.space_before = Pt(8)
    _add_run(missing_heading, "距离全字段 100% 仍缺少的内容", 9.5, True, RGBColor(156, 92, 38))
    incomplete_items = tuple(
        item for item in (*completion.items, *completion.slots) if item.status != "complete"
    )
    if incomplete_items:
        state_labels = {
            "partial": "部分完成", "conflict": "冲突",
            "later_update": "后续公开更新", "missing": "缺失",
        }
        reason_labels = {
            "partial": "现有证据仅部分满足要求",
            "conflict": "冲突待解决",
            "later_update": "晚于项目资料截止日，不计入项目期完整度",
            "missing": "资料缺失",
        }
        for state in ("partial", "conflict", "later_update", "missing"):
            grouped = [item for item in incomplete_items if item.status == state]
            if not grouped:
                continue
            group_heading = doc.add_paragraph()
            group_heading.paragraph_format.space_before = Pt(5)
            _add_run(
                group_heading, f"{state_labels[state]}（{len(grouped)}）", 9, True,
                RGBColor(156, 92, 38),
            )
            for item in grouped:
                paragraph = doc.add_paragraph(style="List Bullet")
                _add_run(
                    paragraph, f"{item.label}（{item.key}）：{reason_labels[state]}", 8.5,
                )
    else:
        doc.add_paragraph("无：全字段覆盖率已达到 100%。")

    prompt = doc.add_paragraph()
    prompt.paragraph_format.space_before = Pt(10)
    _add_run(prompt, "请作者确认：本稿是否完稿？", 10.5, True, RGBColor(21, 40, 64))
    choice = doc.add_paragraph()
    _add_run(choice, "□ 已完稿    □ 尚未完稿（继续补充/修订）", 10, False)
    recommendation = doc.add_paragraph()
    ready = (
        completion.required_rate == 100.0
        and completion.required_slot_rate == 100.0
        and formula_status != "fail"
    )
    text = (
        "系统建议：核心资料已齐，可进入作者终审；独立验证率和公式门禁仍应结合上表判断。"
        if ready else
        "系统建议：尚未达到核心资料或公式门禁要求，当前应作为未完稿继续补充。"
    )
    _add_run(recommendation, text, 8.5, False, RGBColor(105, 113, 122))


def _bullet_block(doc, facts, schema, snapshot, risk=False):
    for index, (key, fact) in enumerate(facts, start=1):
        paragraph = doc.add_paragraph(style="List Bullet")
        if risk:
            _add_run(paragraph, f"风险 {index}：", 9.5, True, RGBColor(156, 92, 38))
        _add_run(paragraph, _fact_value(key, fact, schema, snapshot), 9.5)
        source = doc.add_paragraph()
        source.paragraph_format.left_indent = Inches(0.25)
        source.paragraph_format.space_after = Pt(7)
        _add_run(source, f"来源：{_fact_sources(key, fact, snapshot)}", 7.5, False, RGBColor(105, 113, 122))


def _standard_fact_table(doc, facts, schema, snapshot, include_uri=False, source_aliases=None):
    headers = ("指标", "数值", "证据状态", "来源位置") if include_uri else ("指标", "数值", "证据状态")
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    _repeat_header(table.rows[0])
    for cell, text in zip(table.rows[0].cells, headers):
        _shade_cell(cell, "1F4E78")
        _add_run(cell.paragraphs[0], text, 9, True, RGBColor(255, 255, 255))
    for key, fact in facts:
        row = table.add_row()
        _keep_row_together(row)
        cells = row.cells
        rendered_value = _fact_value(key, fact, schema, snapshot)
        if include_uri and fact.status != "conflict":
            rendered_value = _preview(rendered_value)
        values = [
            _field_for_key(schema, key).label,
            rendered_value,
            _evidence_status(fact),
        ]
        if include_uri:
            values.append(_fact_sources(
                key, fact, snapshot, include_uri=include_uri, source_aliases=source_aliases,
            ))
        for cell, value in zip(cells, values):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            _add_run(cell.paragraphs[0], str(value), 8)
        for cell in cells[1:]:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if not include_uri:
        _body_fact_sources(doc, facts, snapshot)


def _series_fact_table(doc, facts, schema, snapshot):
    grouped = {}
    periods = []
    for key, fact in facts:
        field = _field_for_key(schema, key)
        period = _period(key)
        if not period:
            continue
        grouped.setdefault(field.key, {})[period] = (key, fact)
        if period not in periods:
            periods.append(period)
    periods.sort(key=_period_sort_key)
    if not grouped:
        _standard_fact_table(doc, facts, schema, snapshot)
        return
    table = doc.add_table(rows=1, cols=len(periods) + 1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for cell, value in zip(table.rows[0].cells, ("指标", *periods)):
        _shade_cell(cell, "1F4E78")
        _add_run(cell.paragraphs[0], value, 9, True, RGBColor(255, 255, 255))
    for field_key in sorted(grouped):
        cells = table.add_row().cells
        field = next(field for field in schema.fields if field.key == field_key)
        values = [field.label]
        for period in periods:
            item = grouped[field_key].get(period)
            if item is None:
                values.append("—")
                continue
            key, fact = item
            values.append(_fact_value(key, fact, schema, snapshot))
        for cell, value in zip(cells, values):
            _add_run(cell.paragraphs[0], str(value), 8)
        for cell in cells[1:]:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _body_fact_sources(doc, facts, snapshot)


def _body_fact_sources(doc, facts, snapshot):
    records = []
    for key, fact in facts:
        records.extend(snapshot.candidates.get(key, (fact,)) if fact.status == "conflict" else (fact,))
    if not records:
        return
    by_uri = {}
    for record in records:
        by_uri.setdefault(record.source_uri, []).append(record)
    parts = [
        f"{_source_name(uri)}#{_compact_source_records(group)}"
        for uri, group in sorted(by_uri.items(), key=lambda item: str(item[0]).casefold())
    ]
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(6)
    _add_run(paragraph, f"模型依据：{'；'.join(parts)}", 7.5, False, RGBColor(105, 113, 122))


def _period(key):
    match = re.search(r"\.((?:19|20)\d{2}(?:[aAeE]|_\d{1,2}[mM]))$", key)
    return match.group(1).upper() if match else ""


def _period_sort_key(value):
    year = int(value[:4])
    interim = re.search(r"_(\d{1,2})M$", value)
    if interim:
        return year, 0, int(interim.group(1))
    suffix = value[-1:]
    return year, {"A": 1, "E": 2}.get(suffix, 1), 0


def _field_for_key(schema, key):
    candidates = [field for field in schema.fields if key == field.key or key.startswith(field.key + ".")]
    if not candidates:
        raise KeyError(f"schema field missing for rendered fact: {key}")
    return max(candidates, key=lambda field: len(field.key))


def _fact_value(key, fact, schema, snapshot):
    candidates = snapshot.candidates.get(key, (fact,))
    rendered = [_display_fact(candidate, _field_for_key(schema, key)) for candidate in candidates]
    if fact.status == "conflict":
        return "冲突：" + "；".join(dict.fromkeys(rendered))
    return rendered[0]


def _source_index_block(doc, facts, schema, snapshot):
    uris = sorted({candidate.source_uri for key, fact in facts
                   for candidate in snapshot.candidates.get(key, (fact,))})
    aliases = {uri: f"S{index}" for index, uri in enumerate(uris, start=1)}
    heading = doc.add_paragraph()
    _add_run(heading, "来源文件索引", 9, True, RGBColor(24, 137, 142))
    for uri in uris:
        paragraph = doc.add_paragraph(style="List Bullet")
        _add_run(paragraph, f"[{aliases[uri]}] {_source_name(uri)}", 8)
    _standard_fact_table(
        doc, facts, schema, snapshot, include_uri=True, source_aliases=aliases,
    )


def _fact_sources(key, fact, snapshot, include_uri=False, source_aliases=None):
    candidates = snapshot.candidates.get(key, (fact,)) if fact.status == "conflict" else (fact,)
    return _compact_source_records(
        candidates, include_uri=include_uri, source_aliases=source_aliases,
    )


def _compact_source_records(records, include_uri=False, source_aliases=None):
    groups, raw = {}, []
    for record in records:
        match = re.fullmatch(r"(.+)!([A-Za-z]+)(\d+)", record.source_location)
        prefix = ((f"[{source_aliases[record.source_uri]}]" if source_aliases else _source_name(record.source_uri))
                  if include_uri else "")
        if not match:
            raw.append(f"{prefix}#{record.source_location}" if prefix else record.source_location)
            continue
        sheet, column, row = match.groups()
        groups.setdefault((prefix, sheet, int(row)), set()).add(column_index_from_string(column))
    parts = []
    for (prefix, sheet, row), columns in groups.items():
        coordinates = _compact_columns(sorted(columns), row)
        location = f"{sheet}!{coordinates}"
        parts.append(f"{prefix}#{location}" if prefix else location)
    parts.extend(raw)
    return "；".join(dict.fromkeys(parts))


def _compact_columns(columns, row):
    ranges = []
    start = previous = columns[0]
    for column in columns[1:] + [None]:
        if column is not None and column == previous + 1:
            previous = column
            continue
        first = f"{get_column_letter(start)}{row}"
        last = f"{get_column_letter(previous)}{row}"
        ranges.append(first if start == previous else f"{first}:{last}")
        if column is not None:
            start = previous = column
    return ",".join(ranges)


def _source_name(uri):
    return uri if str(uri).startswith(("http://", "https://")) else Path(uri).name


def _preview(value, limit=56):
    text = str(value).replace("\n", " ").strip()
    return text if len(text) <= limit else text[:limit - 1].rstrip() + "…"


def _display_fact(fact, field):
    value, unit = fact.value, fact.unit or ""
    if isinstance(value, (int, float)) and not isinstance(value, bool):
        number = float(value)
        if unit == "%" or field.unit_family == "percent":
            number = number * 100 if abs(number) <= 1.5 else number
            return f"{_number(number, 2)}%"
        if unit.casefold() == "x" or field.unit_family == "multiple":
            return f"{_number(number, 2)}x"
        rendered = _number(number, 4)
        if not unit and field.unit_family == "currency":
            unit = "model unit"
        return f"{rendered}{unit}"
    return f"{value}{unit}"


def _evidence_status(fact):
    """Evidence status label for the memo table; flags formula-derived facts."""
    if fact.status == "conflict":
        return "冲突"
    if fact.status == "missing":
        return "缺失"
    if fact.status == "retained_original":
        return "原文保留"
    # 公式推导事实: evidence 中带 formula= 前缀 (discover 的 rule 是匹配规则名,
    # 公式信息在证据链里; 值已经过 LibreOffice 重算覆盖, 标注"公式重算").
    if any(str(line).startswith("formula=") for line in fact.evidence):
        return "公式重算"
    if fact.status == "verified":
        return "已验证"
    if fact.status == "corroborated":
        return "交叉验证"
    return "来源真值"


def _number(value, decimals):
    rendered = f"{value:,.{decimals}f}"
    return rendered.rstrip("0").rstrip(".")


def _add_run(paragraph, text, size, bold=False, color=None, east=None):
    run = paragraph.add_run(text)
    run.font.name = LATIN_FONT
    run._element.get_or_add_rPr().rFonts.set(
        qn("w:eastAsia"), east or EAST_ASIA_FONT)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def _fact_table(doc, snapshot, prefixes):
    facts = [(key, fact) for key, fact in sorted(snapshot.selected.items()) if key.startswith(prefixes)]
    if not facts:
        doc.add_paragraph("No mapped facts were supplied for this section.")
        return
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for cell, text in zip(table.rows[0].cells, ("Metric", "Value", "Status", "Source")):
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        _add_run(cell.paragraphs[0], text, 9, True)
    for key, fact in facts:
        cells = table.add_row().cells
        value = f"{_display_value(fact.value)} {fact.unit or 'model unit'}"
        for cell, text in zip(cells, (key, value, fact.status, fact.source_location)):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            _add_run(cell.paragraphs[0], str(text), 8.5)


def _display_value(value):
    if isinstance(value, float):
        return f"{value:.4f}".rstrip("0").rstrip(".")
    return str(value)
